import sys
import os
import docx


DIR = "../КР 9"
TEMPLATE_DOC = None

def get_style_files(fname):

    docfile = docx.Document(fname)

    if len(docfile.paragraphs) < 1:
        return None

    return docfile.paragraphs[0].runs[0].style

def format_folder_style(path, styl):

    for root, dirs, files in os.walk(path):
        for fname in files:
            if not fname.endswith(".docx"):
                continue

            newdoc = docx.Document(fname)
            for par in newdoc.paragraphs:
                for run in par.runs:
                    run.style = styl

            newdoc.save("dfgdghdh.docx")


if __name__ == '__main__':
    if len(sys.argv) < 3:
        DIR = input("Folder: ")
        TEMPLATE_DOC = input("Template file: ")
        if TEMPLATE_DOC is None or not TEMPLATE_DOC.endswith(".docx"):
            print(f"Invalid file name {TEMPLATE_DOC}")
            exit()

        path_template = os.path.join(DIR, TEMPLATE_DOC)

        if not os.path.exists(path_template) or not os.path.isfile(path_template):
            print(f"Invalid file name {path_template}")
            exit()

        st = get_style_files(path_template)
        print(st)
        format_folder_style(DIR, st)
